import os.path
from os import listdir
from dotenv import load_dotenv
from os.path import isfile, join
from typing import Literal, get_args
from langchain.chains import RetrievalQA
from langchain.schema.document import Document
from langchain_community.tools import WikipediaQueryRun
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.utilities import WikipediaAPIWrapper 
from langchain.text_splitter import RecursiveCharacterTextSplitter 
from langchain_community.vectorstores import DocArrayInMemorySearch
import docx

DataSource = Literal["Wikipedia", "Documents"]
SUPPORTED_DATA_SOURCES = get_args(DataSource)

# loading API keys from env
load_dotenv()
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")

if OPENAI_API_KEY == 'xxxxxxxx':
    raise ValueError("Please add your own OpenAI API key in the .env file by replacing 'xxxxxxxx' with your own key.")

# loading model and defining embedding
llm = ChatOpenAI(temperature=0, model='gpt-3.5-turbo')
embeddings = OpenAIEmbeddings()

# get target folder for uploaded docs
target_folder = "./docs/"

def load_data_set(source: DataSource, query: str):
    if source not in SUPPORTED_DATA_SOURCES:
        raise ValueError(f"Provided data source {source} is not supported.")

    # fragmenting the document content to fit in the number of token limitations
    text_splitter = RecursiveCharacterTextSplitter(chunk_size = 500, chunk_overlap = 50)

    if source == "Wikipedia":
        Wikipedia = WikipediaQueryRun(api_wrapper=WikipediaAPIWrapper())
        data = Wikipedia.run(query)
        split_docs = [Document(page_content=sent) for sent in data.split('\n')]
    elif source == "Documents":
    # get files from target directory
        my_file = [f for f in listdir(target_folder) if isfile(join(target_folder, f))][0]
        my_file = target_folder + my_file
        print(f"my file is {my_file}")

        if my_file.endswith('.pdf'):
            # load uploaded pdf file
            loader = PyPDFLoader(my_file)
            data = loader.load()
            split_docs = text_splitter.split_documents(data)
        elif my_file.endswith('.docx'):
            # load uploaded docx file
            doc = docx.Document(my_file)

            # parse data from docx file
            data_parsed = ''
            for para in doc.paragraphs:
                data_parsed += para.text
            
            # Create a Document object from the parsed data
            doc_object = Document(page_content=data_parsed)
            split_docs = text_splitter.split_documents([doc_object])  # Pass a list of Document objects

    data_set = DocArrayInMemorySearch.from_documents(documents = split_docs, embedding = embeddings)

    return data_set


def retrieve_info(source: DataSource, data_set: DocArrayInMemorySearch, query: str):
    if source not in SUPPORTED_DATA_SOURCES:
        raise ValueError(f"Provided data source {source} is not supported.")

    qa = RetrievalQA.from_chain_type(
        llm = llm,
        chain_type="stuff",
        retriever = data_set.as_retriever(), # repla
        verbose=True,
    )

    output = qa.invoke(query)

    return output


def generate_answer(selection: DataSource, query: str):
    if selection not in SUPPORTED_DATA_SOURCES:
        raise ValueError(f"Provided data source {selection} is not supported.")

    data_set = load_data_set(selection, query)
    response = retrieve_info(selection, data_set, query)
    
    return response